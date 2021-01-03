#!/usr/bin/env python
from docx import Document

# file name
word_write_file_name = "word.docx"

# create document object
# <class 'docx.document.Document'>
document = Document()

# write text
document.add_paragraph("hellow word 1 page.")
document.paragraphs[0].add_run("add text")

# add new page
document.add_page_break()

# write text
document.add_paragraph("Hello Word 2 page.")
document.add_paragraph("Hello Word 3 page?")

# write it into word file
document.save(word_write_file_name)